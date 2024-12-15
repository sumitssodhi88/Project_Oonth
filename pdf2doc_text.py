import os
import pdfplumber
from tkinter import Tk, Label, StringVar, Frame
from tkinter.filedialog import askdirectory
import threading
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def update_status(status_var, text):
    status_var.set(text)
    root.update_idletasks()


def should_end_line(word):
    if word[-1] in [".", ";", "]"]:  
        return True
    if word.isdigit():  
        return True
    if word[0].isupper():  
        return True
    return False


def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run._r.append(OxmlElement('w:br'))
    run._r[-1].set(qn('w:type'), 'page')


def process_pdfs():
    folder_path = askdirectory(title="Select the folder containing your PDFs")

    if not folder_path: 
        update_status(status_var, "Process canceled.")
        root.after(2000, root.quit)
        return

    output_folder = os.path.join(folder_path, "pdf_to_docs_sectionless")
    os.makedirs(output_folder, exist_ok=True)

    pdf_files = [f for f in os.listdir(folder_path) if f.lower().endswith(".pdf")]
    total_files = len(pdf_files)

    for pdf_index, filename in enumerate(pdf_files, start=1):
        pdf_path = os.path.join(folder_path, filename)
        doc_filename = os.path.splitext(filename)[0] + ".docx"
        doc_path = os.path.join(output_folder, doc_filename)

        update_status(status_var, f"Processing {filename} ({pdf_index}/{total_files})...")

        with pdfplumber.open(pdf_path) as pdf:
            doc = Document()

            total_pages = len(pdf.pages)

            previous_line = ""  
            for page_index, page in enumerate(pdf.pages, start=1):
                page_text = page.extract_text(x_tolerance=2, y_tolerance=2)
                if page_text:
                    lines = page_text.split('\n')

                    for line_index, line in enumerate(lines):
                        words = line.split()

                        if previous_line:
                            words.insert(0, previous_line)
                            previous_line = ""  

                        if words and should_end_line(words[-1]):
                            doc.add_paragraph(" ".join(words))
                        else:
                            previous_line = " ".join(words)

                add_page_break(doc)  

                update_status(status_var, f"Processing {filename} - Page {page_index}/{total_pages}")

            doc.save(doc_path)

        update_status(status_var, f"Finished {filename} ({pdf_index}/{total_files})")

    update_status(status_var, "Conversion complete! Word files saved in the 'converted_docs' folder.")
    root.after(2000, root.quit)


root = Tk()
root.title("PDF to DOCX Conversion")

frame = Frame(root)
frame.pack(pady=10)

status_var = StringVar()
status_var.set("Ready to process...")

status_label = Label(frame, textvariable=status_var, anchor="w", width=50, height=2, bg="lightgrey")
status_label.pack()

try:
    thread = threading.Thread(target=process_pdfs)
    thread.start()
except Exception as e:
    update_status(status_var, f"Error starting thread: {e}")
    root.quit()

root.mainloop()
